#!/usr/bin/env python3
"""office-docs helper — build Word .docx and Excel .xlsx files.

Usage:  python office.py '<json-spec>'   (or pipe the JSON on stdin)
Ops:
  docx {out, title?, blocks:[{type:heading|para|bullets|table, ...}]}  -> {out, blocks}
  xlsx {out, sheets:{name:[[row]]}} | {out, rows:[[...]], sheet?}       -> {out, sheets}

A fast start, not a cage: for styles, images, headers/footers, or cell formatting,
use python-docx / openpyxl directly.
"""
import json
import sys


def read_spec():
    if len(sys.argv) > 1 and sys.argv[1].strip():
        return json.loads(sys.argv[1])
    data = sys.stdin.read()
    if data.strip():
        return json.loads(data)
    raise ValueError("no spec: pass a JSON spec as argv[1] or on stdin")


def op_docx(spec):
    from docx import Document

    doc = Document()
    if spec.get("title"):
        doc.add_heading(str(spec["title"]), level=0)
    blocks = spec.get("blocks") or []
    for b in blocks:
        kind = b.get("type")
        if kind == "heading":
            doc.add_heading(str(b.get("text", "")), level=int(b.get("level", 1)))
        elif kind == "para":
            doc.add_paragraph(str(b.get("text", "")))
        elif kind == "bullets":
            for item in b.get("items") or []:
                doc.add_paragraph(str(item), style="List Bullet")
        elif kind == "table":
            header = b.get("header") or []
            rows = b.get("rows") or []
            ncols = len(header) if header else (len(rows[0]) if rows else 0)
            if ncols:
                t = doc.add_table(rows=0, cols=ncols)
                t.style = "Light Grid Accent 1"
                if header:
                    cells = t.add_row().cells
                    for i, h in enumerate(header):
                        cells[i].text = str(h)
                for r in rows:
                    cells = t.add_row().cells
                    for i in range(ncols):
                        cells[i].text = str(r[i]) if i < len(r) else ""
        else:
            raise ValueError(f"unknown block type: {kind}")
    out = spec.get("out", "document.docx")
    doc.save(out)
    return {"out": out, "blocks": len(blocks)}


def op_xlsx(spec):
    from openpyxl import Workbook
    from openpyxl.styles import Font

    sheets = spec.get("sheets")
    if not sheets:
        rows = spec.get("rows")
        if rows is None:
            raise ValueError("xlsx needs sheets{} or rows[]")
        sheets = {spec.get("sheet", "Sheet1"): rows}

    wb = Workbook()
    wb.remove(wb.active)  # drop the default sheet; we add our own
    for name, rows in sheets.items():
        ws = wb.create_sheet(title=str(name)[:31])  # Excel caps sheet names at 31
        for r in rows or []:
            ws.append(list(r))
        if rows:
            for cell in ws[1]:  # bold the header row
                cell.font = Font(bold=True)
            ws.freeze_panes = "A2"
    out = spec.get("out", "workbook.xlsx")
    wb.save(out)
    return {"out": out, "sheets": list(sheets.keys())}


OPS = {"docx": op_docx, "xlsx": op_xlsx}


def run(spec):
    op = spec.get("op")
    if op not in OPS:
        raise ValueError("spec.op must be one of: " + ", ".join(OPS))
    result = OPS[op](spec)
    result.update({"ok": True, "op": op})
    return result


def main():
    try:
        print(json.dumps(run(read_spec()), default=str))
    except Exception as e:  # noqa: BLE001
        print(json.dumps({"ok": False, "error": str(e)}))
        sys.exit(1)


if __name__ == "__main__":
    main()
